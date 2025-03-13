from flask import Flask, render_template, request, redirect, url_for, send_file, make_response, send_from_directory, flash
from flask import Flask, render_template, request, redirect, url_for, send_file, make_response, send_from_directory, flash
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from datetime import datetime, timedelta
from io import BytesIO
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pyngrok import ngrok
import os
from functools import wraps
from flask import request, Response
import urllib.parse

app = Flask(__name__)

# Configuração do SQLite para a Vercel
import os
database_url = os.environ.get('DATABASE_URL', 'sqlite:///pedidos.db')

if database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'A8Zyr47j3yX99Kq2QdNf'

db = SQLAlchemy(app)
migrate = Migrate(app, db)

class Pedido(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome_aluno = db.Column(db.String(100), nullable=False)
    whatsapp = db.Column(db.String(20), nullable=False)
    tamanho_camisa = db.Column(db.String(10), nullable=False)
    modelo_camisa = db.Column(db.String(20), nullable=False)
    quer_moletom = db.Column(db.Boolean, default=False)
    modelo_moletom = db.Column(db.String(20), nullable=True)
    tamanho_moletom = db.Column(db.String(10), nullable=True)
    nome_na_camisa = db.Column(db.String(100), nullable=False)
    pago = db.Column(db.Boolean, default=False)
    data_pedido = db.Column(db.DateTime, default=datetime.utcnow)
    primeira_parcela_paga = db.Column(db.Boolean, default=False)
    segunda_parcela_paga = db.Column(db.Boolean, default=False)
    data_primeira_parcela = db.Column(db.DateTime, nullable=True)
    data_segunda_parcela = db.Column(db.DateTime, nullable=True)

with app.app_context():
    db.create_all()

def check_auth(username, password):
    """Verifica se o username e senha estão corretos."""
    return username == 'admin' and password == 'admin123'

def authenticate():
    """Envia um header de autenticação básica."""
    return Response(
        'Área restrita. Por favor, faça login.\n',
        401,
        {'WWW-Authenticate': 'Basic realm="Login Required"'}
    )

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

@app.route('/')
def index():
    return render_template('formulario.html')

@app.route('/criar_pedido', methods=['POST'])
def criar_pedido():
    print("Dados recebidos do formulário:", request.form)  # Debug: imprime os dados do formulário
    quer_moletom = 'quer_moletom' in request.form
    
    modelo_moletom = request.form.get('modelo_moletom') if quer_moletom else None
    tamanho_moletom = request.form.get('tamanho_moletom') if quer_moletom else None

    novo_pedido = Pedido(
        nome_aluno=request.form.get('nome_aluno'),
        whatsapp=request.form.get('whatsapp'),
        tamanho_camisa=request.form.get('tamanho_camisa'),
        modelo_camisa=request.form.get('modelo_camisa'),
        quer_moletom=quer_moletom,
        modelo_moletom=modelo_moletom,
        tamanho_moletom=tamanho_moletom,
        nome_na_camisa=request.form.get('nome_na_camisa'),
        data_pedido=datetime.now()
    )
    
    try:
        db.session.add(novo_pedido)
        db.session.commit()
        flash('Pedido enviado com sucesso!')
        return redirect(url_for('index'))
    except Exception as e:
        db.session.rollback()
        print("Erro ao criar pedido:", str(e))  # Debug: imprime o erro
        flash('Erro ao criar pedido: ' + str(e))
        return redirect(url_for('index'))

@app.route('/pedidos')
@requires_auth
def lista_pedidos():
    pedidos = Pedido.query.all()
    # Debug: imprimir informações dos pedidos
    for pedido in pedidos:
        print(f"""
        Pedido ID: {pedido.id}
        Nome: {pedido.nome_aluno}
        Quer Moletom: {pedido.quer_moletom}
        Modelo Moletom: {pedido.modelo_moletom}
        Tamanho Moletom: {pedido.tamanho_moletom}
        """)
    return render_template('pedidos.html', pedidos=pedidos)

@app.route('/atualizar_pagamento/<int:id>')
def atualizar_pagamento(id):
    pedido = Pedido.query.get_or_404(id)
    pedido.pago = not pedido.pago
    db.session.commit()
    return redirect(url_for('lista_pedidos'))

@app.route('/deletar_pedido/<int:id>')
def deletar_pedido(id):
    pedido = Pedido.query.get_or_404(id)
    db.session.delete(pedido)
    db.session.commit()
    return redirect(url_for('lista_pedidos'))

@app.route('/exportar/<formato>')
def exportar(formato):
    pedidos = Pedido.query.all()
    
    if formato == 'excel':
        # Criar DataFrame com pandas
        dados = []
        for p in pedidos:
            moletom_info = f"{p.modelo_moletom} - {p.tamanho_moletom}" if p.quer_moletom else "Não solicitado"
            dados.append({
                'Nome': p.nome_aluno,
                'WhatsApp': p.whatsapp,
                'Modelo Camisa': p.modelo_camisa,
                'Tamanho Camisa': p.tamanho_camisa,
                'Moletom': moletom_info,
                'Nome na Camisa': p.nome_na_camisa,
                'Status 1ª Parcela': 'Paga' if p.primeira_parcela_paga else 'Pendente',
                'Data 1ª Parcela': p.data_primeira_parcela.strftime('%d/%m/%Y') if p.data_primeira_parcela else '-',
                'Status 2ª Parcela': 'Paga' if p.segunda_parcela_paga else 'Pendente',
                'Data 2ª Parcela': p.data_segunda_parcela.strftime('%d/%m/%Y') if p.data_segunda_parcela else '-',
                'Data do Pedido': p.data_pedido.strftime('%d/%m/%Y %H:%M')
            })
        
        df = pd.DataFrame(dados)
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Pedidos')
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='pedidos.xlsx'
        )
    
    elif formato == 'word':
        try:
            doc = Document()
            doc.add_heading('Lista de Pedidos', 0)
            
            # Adiciona tabela
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Define larguras das colunas (em polegadas)
            widths = [1.5, 1.5, 2.0, 2.0, 1.5, 2.0]
            for idx, width in enumerate(widths):
                for cell in table.columns[idx].cells:
                    cell.width = width * 914400  # Converte polegadas para twips
            
            # Cabeçalho
            header_cells = table.rows[0].cells
            headers = ['Nome', 'WhatsApp', 'Camisa', 'Moletom', 'Nome na Camisa']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Dados
            for pedido in pedidos:
                row_cells = table.add_row().cells
                
                # Nome
                row_cells[0].text = pedido.nome_aluno
                
                # WhatsApp
                row_cells[1].text = pedido.whatsapp
                
                # Camisa
                row_cells[2].text = f"Modelo: {pedido.modelo_camisa}\nTamanho: {pedido.tamanho_camisa}"
                
                # Moletom
                if pedido.quer_moletom:
                    row_cells[3].text = f"Modelo: {pedido.modelo_moletom}\nTamanho: {pedido.tamanho_moletom}"
                else:
                    row_cells[3].text = "Não solicitado"
                
                # Nome na Camisa
                row_cells[4].text = pedido.nome_na_camisa
                
              
            # Ajusta o estilo da tabela
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.style = doc.styles['Normal']
                        paragraph.paragraph_format.space_before = 0
                        paragraph.paragraph_format.space_after = 0
                        for run in paragraph.runs:
                            run.font.size = 127000  # Tamanho 10pt
            
            # Salva o documento
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='pedidos.docx'
            )
            
        except Exception as e:
            print(f"Erro ao gerar DOCX: {str(e)}")  # Log do erro
            flash('Erro ao gerar arquivo Word. Por favor, tente novamente.', 'error')
            return redirect(url_for('lista_pedidos'))
    
    elif formato == 'pdf':
        try:
            output = BytesIO()
            c = canvas.Canvas(output, pagesize=A4)
            
            # Título
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 800, "Lista de Pedidos")
            
            # Configurações iniciais
            y = 750
            c.setFont("Helvetica-Bold", 10)
            
            # Cabeçalho
            headers = ['Nome', 'WhatsApp', 'Camisa', 'Moletom', 'Nome na Camisa']
            x_positions = [50, 150, 250, 350, 450, 550]
            
            # Desenha linha superior
            c.line(45, y+15, 550, y+15)
            
            # Desenha cabeçalhos
            for header, x in zip(headers, x_positions):
                c.drawString(x, y, header)
            
            # Linha após cabeçalhos
            c.line(45, y-5, 550, y-5)
            
            # Conteúdo
            y -= 30
            c.setFont("Helvetica", 10)
            
            for pedido in pedidos:
                # Verifica espaço na página
                if y < 50:
                    c.showPage()
                    y = 750
                    c.setFont("Helvetica", 10)
                
                # Nome
                c.drawString(50, y, pedido.nome_aluno[:20])
                
                # WhatsApp
                c.drawString(150, y, pedido.whatsapp)
                
                # Camisa
                c.drawString(250, y, f"Modelo: {pedido.modelo_camisa}")
                c.drawString(250, y-12, f"Tam: {pedido.tamanho_camisa}")
                
                # Moletom
                if pedido.quer_moletom:
                    c.drawString(350, y, f"Modelo: {pedido.modelo_moletom}")
                    c.drawString(350, y-12, f"Tam: {pedido.tamanho_moletom}")
                else:
                    c.drawString(350, y, "Não solicitado")
                
                # Nome na Camisa
                c.drawString(450, y, pedido.nome_na_camisa[:15])
                
               
                
                # Linha separadora
                y -= 35
                c.line(45, y+5, 550, y+5)
                y -= 10
            
            c.save()
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='pedidos.pdf'
            )
        
        except Exception as e:
            print(f"Erro ao gerar PDF: {str(e)}")  # Log do erro
            flash('Erro ao gerar PDF. Por favor, tente novamente.', 'error')
            return redirect(url_for('lista_pedidos'))

@app.route('/enviar_whatsapp/<int:id>/<parcela>')
@requires_auth
def enviar_whatsapp(id, parcela):
    pedido = Pedido.query.get_or_404(id)
    valor_total = "295,00" if pedido.quer_moletom else "95,00"
    valor_parcela = "147,50" if pedido.quer_moletom else "47,50"
    
    mensagem = f"Olá {pedido.nome_aluno}, sobre seu pedido de uniforme:\n\n"
    mensagem += f"Valor total: R$ {valor_total}\n"
    mensagem += f"Valor da parcela: R$ {valor_parcela}\n\n"
    mensagem += f"PIX: jamersonpontes25@gmail.com\n"
    mensagem += f"Enviar comprovante após o pagamento.\n\n"
    
    if parcela == "primeira":
        mensagem += "Esta é a cobrança da PRIMEIRA parcela."
    else:
        mensagem += "Esta é a cobrança da SEGUNDA parcela."
    
    # Formata o número do WhatsApp (remove caracteres não numéricos)
    whatsapp = ''.join(filter(str.isdigit, pedido.whatsapp))
    
    url = f"https://wa.me/55{whatsapp}?text={urllib.parse.quote(mensagem)}"
    return redirect(url)

@app.route('/static/sw.js')
def sw():
    response = make_response(send_from_directory('static', 'sw.js'))
    response.headers['Content-Type'] = 'application/javascript'
    return response

@app.route('/marcar_primeira_parcela/<int:id>')
@requires_auth
def marcar_primeira_parcela(id):
    pedido = Pedido.query.get_or_404(id)
    pedido.primeira_parcela_paga = True
    pedido.data_primeira_parcela = datetime.now()
    db.session.commit()
    flash('Primeira parcela marcada como paga!')
    return redirect(url_for('lista_pedidos'))

@app.route('/marcar_segunda_parcela/<int:id>')
@requires_auth
def marcar_segunda_parcela(id):
    pedido = Pedido.query.get_or_404(id)
    pedido.segunda_parcela_paga = True
    pedido.data_segunda_parcela = datetime.now()
    db.session.commit()
    flash('Segunda parcela marcada como paga!')
    return redirect(url_for('lista_pedidos'))

@app.route('/cancelar_pagamento/<int:id>')
@requires_auth
def cancelar_pagamento(id):
    pedido = Pedido.query.get_or_404(id)
    pedido.primeira_parcela_paga = False
    pedido.segunda_parcela_paga = False
    pedido.data_primeira_parcela = None
    pedido.data_segunda_parcela = None
    db.session.commit()
    flash('Pagamentos cancelados!')
    return redirect(url_for('lista_pedidos'))

# Criar todas as tabelas
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000))) 